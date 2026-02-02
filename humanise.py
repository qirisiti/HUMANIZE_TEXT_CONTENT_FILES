import os
import sys
import requests
from pathlib import Path
from dotenv import load_dotenv
import anthropic
from docx import Document
import PyPDF2
import re
import copy

# Load environment variables
load_dotenv()

ZEROGPT_API_KEY = os.getenv("ZEROGPT_API_KEY")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

BASE_DIR = Path(__file__).parent
FILES_DIR = BASE_DIR / "files_to_check"
OUTPUT_DIR = BASE_DIR / "output"

ZEROGPT_DETECT_FILE_URL = "https://api.zerogpt.com/api/detect/detectFile"
ZEROGPT_DETECT_TEXT_URL = "https://api.zerogpt.com/api/detect/detectText"

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt"}

# Patterns to detect citations
CITATION_PATTERNS = [
    r'\([A-Z][a-zA-Z]+(?:\s+(?:et\s+al\.?|&|and)\s+[A-Z][a-zA-Z]+)*,?\s*\d{4}[a-z]?\)',  # (Author, 2023) or (Author et al., 2023)
    r'\([A-Z][a-zA-Z]+\s+&\s+[A-Z][a-zA-Z]+,?\s*\d{4}\)',  # (Author & Author, 2023)
    r'\[\d+(?:[-,]\s*\d+)*\]',  # [1] or [1,2] or [1-3]
    r'\([A-Z][a-zA-Z]+,\s*\d{4},\s*p\.?\s*\d+\)',  # (Author, 2023, p. 45)
]

# Keywords that indicate references section
REFERENCES_KEYWORDS = [
    'references', 'bibliography', 'works cited', 'citations', 'sources'
]

# Keywords that indicate title page elements
TITLE_PAGE_KEYWORDS = [
    'submitted to', 'submitted by', 'student id', 'student number', 'registration',
    'course code', 'module', 'supervisor', 'lecturer', 'professor', 'university',
    'college', 'faculty', 'department', 'school of', 'bachelor', 'master',
    'assignment', 'thesis', 'dissertation', 'project', 'date:', 'name:'
]


def is_heading(text: str) -> bool:
    """Check if a line is likely a heading/title."""
    text = text.strip()
    if not text:
        return False

    # Short lines that are likely headings
    if len(text) < 100:
        # Numbered headings (1. Introduction, 1.1 Background)
        if re.match(r'^\d+(\.\d+)*\.?\s+\w', text):
            return True
        # ALL CAPS headings
        if text.isupper() and len(text.split()) <= 8:
            return True
        # Title Case headings (common patterns)
        words = text.split()
        if len(words) <= 8 and text.istitle():
            return True
        # Chapter/Section headings
        if re.match(r'^(chapter|section|part|appendix)\s+\w', text, re.IGNORECASE):
            return True

    return False


def is_title_page_content(text: str) -> bool:
    """Check if text is part of title page."""
    text_lower = text.lower()
    for keyword in TITLE_PAGE_KEYWORDS:
        if keyword in text_lower:
            return True
    return False


def is_references_section(text: str) -> bool:
    """Check if we've entered the references section."""
    text_lower = text.strip().lower()
    for keyword in REFERENCES_KEYWORDS:
        if text_lower == keyword or text_lower.startswith(keyword + '\n'):
            return True
        if re.match(rf'^{keyword}\s*$', text_lower):
            return True
    return False


def protect_citations(text: str) -> tuple[str, dict]:
    """Replace citations with placeholders to protect them."""
    placeholders = {}
    protected_text = text
    counter = 0

    for pattern in CITATION_PATTERNS:
        matches = re.finditer(pattern, protected_text)
        for match in matches:
            citation = match.group()
            placeholder = f"__CITATION_{counter}__"
            placeholders[placeholder] = citation
            protected_text = protected_text.replace(citation, placeholder, 1)
            counter += 1

    return protected_text, placeholders


def restore_citations(text: str, placeholders: dict) -> str:
    """Restore citations from placeholders."""
    restored_text = text
    for placeholder, citation in placeholders.items():
        restored_text = restored_text.replace(placeholder, citation)
    return restored_text


def check_api_keys():
    """Verify API keys are set."""
    if not ZEROGPT_API_KEY:
        print("Error: ZEROGPT_API_KEY not set in .env file")
        sys.exit(1)
    if not ANTHROPIC_API_KEY:
        print("Error: ANTHROPIC_API_KEY not set in .env file")
        sys.exit(1)


def list_files():
    """List all supported files in the files_to_check directory."""
    files = []
    for f in FILES_DIR.iterdir():
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(f)
    return sorted(files)


def detect_ai_in_file(file_path: Path) -> dict:
    """Send file to ZeroGPT for AI detection."""
    headers = {"ApiKey": ZEROGPT_API_KEY}

    try:
        with open(file_path, "rb") as f:
            files = {"file": (file_path.name, f)}
            response = requests.post(ZEROGPT_DETECT_FILE_URL, headers=headers, files=files, timeout=60)

        if response.status_code == 200:
            result = response.json()
            # Debug output
            if not result.get("success"):
                print(f"  ZeroGPT API returned error: {result}")
            return result
        else:
            print(f"  HTTP Error from ZeroGPT: {response.status_code}")
            print(f"  Response: {response.text[:500]}")
            return {"success": False, "message": f"HTTP {response.status_code}"}
    except requests.exceptions.RequestException as e:
        print(f"  Request error: {e}")
        return {"success": False, "message": str(e)}


def detect_ai_in_text(text: str) -> dict:
    """Send text to ZeroGPT for AI detection."""
    headers = {"ApiKey": ZEROGPT_API_KEY}
    payload = {"input_text": text}

    try:
        response = requests.post(ZEROGPT_DETECT_TEXT_URL, headers=headers, json=payload, timeout=60)

        if response.status_code == 200:
            result = response.json()
            if not result.get("success"):
                print(f"  ZeroGPT API returned error: {result}")
            return result
        else:
            print(f"  HTTP Error from ZeroGPT: {response.status_code}")
            print(f"  Response: {response.text[:500]}")
            return {"success": False, "message": f"HTTP {response.status_code}"}
    except requests.exceptions.RequestException as e:
        print(f"  Request error: {e}")
        return {"success": False, "message": str(e)}


def extract_text_from_file(file_path: Path) -> str:
    """Extract text content from a file."""
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif suffix == ".docx":
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    elif suffix == ".pdf":
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text

    return ""


def extract_document_sections(text: str) -> dict:
    """Split document into sections: title_page, body, references."""
    lines = text.split('\n')
    sections = {
        'title_page': [],
        'body': [],
        'references': [],
        'headings': {}  # Maps line index to heading text
    }

    in_title_page = True
    in_references = False
    body_started = False

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Check for references section start
        if is_references_section(stripped):
            in_references = True
            in_title_page = False
            sections['references'].append(line)
            continue

        if in_references:
            sections['references'].append(line)
            continue

        # Check if still in title page area (first few paragraphs with title-like content)
        if in_title_page:
            if is_title_page_content(stripped) or (not body_started and len(sections['title_page']) < 15):
                sections['title_page'].append(line)
                # Check if this looks like a main heading that starts the body
                if is_heading(stripped) and not is_title_page_content(stripped):
                    if any(kw in stripped.lower() for kw in ['introduction', 'abstract', 'chapter 1', '1.']):
                        in_title_page = False
                        body_started = True
                        sections['body'].append(line)
                        sections['headings'][len(sections['body']) - 1] = stripped
                continue
            else:
                in_title_page = False
                body_started = True

        # We're in the body section
        if is_heading(stripped):
            sections['headings'][len(sections['body'])] = stripped

        sections['body'].append(line)

    return sections


def reconstruct_document(sections: dict) -> str:
    """Reconstruct document from sections."""
    parts = []

    if sections['title_page']:
        parts.append('\n'.join(sections['title_page']))

    if sections['body']:
        parts.append('\n'.join(sections['body']))

    if sections['references']:
        parts.append('\n'.join(sections['references']))

    return '\n'.join(parts)


def humanise_text(text: str, ai_sentences: list = None) -> str:
    """Use Claude to humanise the text while keeping formal academic tone."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    # Extract document sections
    sections = extract_document_sections(text)

    # Only humanise the body, protect citations
    body_text = '\n'.join(sections['body'])
    protected_body, citation_placeholders = protect_citations(body_text)

    # Get the headings to pass to Claude
    headings_list = list(sections['headings'].values())
    headings_note = ""
    if headings_list:
        headings_note = "\n\nHEADINGS TO PRESERVE EXACTLY:\n" + "\n".join(f"- {h}" for h in headings_list[:20])

    prompt = f"""You are an expert academic writer helping a BSc graduate student rewrite their work to sound more natural and human-written while maintaining academic standards.

TASK: Rewrite the following text to:
1. Sound naturally human-written (avoid AI-typical patterns)
2. Maintain a formal, academic tone suitable for university assignments
3. Keep the same meaning, facts, and structure
4. Use varied sentence structures and lengths
5. Include natural transitions and flow
6. Avoid overly perfect grammar - include occasional minor stylistic variations that humans naturally have
7. Keep technical terms and subject-specific vocabulary
8. Maintain the same level of formality as a BSc graduate assignment

CRITICAL - DO NOT CHANGE THESE ELEMENTS (copy them EXACTLY as they appear):
1. HEADINGS & SUBHEADINGS: Keep all section titles, chapter names, topic names exactly as written (see list below)
2. PLACEHOLDERS: Keep all __CITATION_X__ placeholders exactly as they appear (these are protected citations)
3. FIGURES & TABLES: Keep figure captions, table titles, and any numbered labels exactly as they are
4. QUOTES: Keep any direct quotations exactly as written (text within quotation marks)
5. TECHNICAL TERMS: Keep specific technical terminology, acronyms, and proper nouns

ONLY REWRITE: The body paragraph text - the actual discussion/content sentences.
{headings_note}

IMPORTANT:
- Do NOT add any preamble or explanation
- Do NOT say "Here is the rewritten text" or similar
- Just output the rewritten text directly
- Keep approximately the same length
- Maintain the exact same document structure and paragraph order

BODY TEXT TO REWRITE:
{protected_body}"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    humanised_body = message.content[0].text

    # Restore citations
    humanised_body = restore_citations(humanised_body, citation_placeholders)

    # Reconstruct full document
    sections['body'] = humanised_body.split('\n')
    full_text = reconstruct_document(sections)

    return full_text


def save_text_to_file(text: str, original_path: Path, version: int) -> Path:
    """Save humanised text to a new file."""
    stem = original_path.stem
    suffix = original_path.suffix.lower()
    new_name = f"{stem}_humanised_v{version}{suffix}"
    output_path = OUTPUT_DIR / new_name

    if suffix == ".txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

    elif suffix == ".docx":
        doc = Document()
        for para in text.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        doc.save(output_path)

    elif suffix == ".pdf":
        # Save as txt since creating PDFs is complex
        output_path = OUTPUT_DIR / f"{stem}_humanised_v{version}.txt"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"  Note: PDF converted to TXT for output")

    return output_path


def parse_ai_percentage(result: dict) -> float:
    """Extract AI percentage from ZeroGPT response."""
    if not result or not result.get("success"):
        return -1

    data = result.get("data") or {}
    fake_percentage = data.get("fakePercentage", 0)

    # Handle if it's a string with % sign
    if isinstance(fake_percentage, str):
        fake_percentage = fake_percentage.replace("%", "").strip()
        try:
            fake_percentage = float(fake_percentage)
        except ValueError:
            fake_percentage = 0

    return float(fake_percentage)


def get_ai_sentences(result: dict) -> list:
    """Extract AI-detected sentences from ZeroGPT response."""
    if not result or not result.get("success"):
        return []

    data = result.get("data") or {}
    sentences = data.get("sentences", [])

    if isinstance(sentences, str):
        # Try to parse if it's a string representation
        return [sentences] if sentences else []

    return sentences if isinstance(sentences, list) else []


def display_results(result: dict, file_name: str):
    """Display detection results."""
    if not result:
        print("  Failed to get results")
        return

    # Debug: print raw response if there's an issue
    if not result.get("success"):
        print(f"  API Error: {result.get('message', 'Unknown error')}")
        print(f"  Raw response: {result}")
        return

    data = result.get("data") or {}
    ai_percentage = parse_ai_percentage(result)
    text_words = data.get("textWords", "N/A")
    ai_words = data.get("aiWords", "N/A")

    print(f"\n  Results for: {file_name}")
    print(f"  {'-' * 40}")
    print(f"  AI Percentage: {ai_percentage}%")
    print(f"  Total Words: {text_words}")
    print(f"  AI Words: {ai_words}")
    print(f"  {'-' * 40}")


def main():
    print("\n" + "=" * 50)
    print("  AI CONTENT HUMANISER")
    print("=" * 50)

    check_api_keys()

    # Ensure directories exist
    FILES_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    # List available files
    files = list_files()

    if not files:
        print(f"\nNo supported files found in: {FILES_DIR}")
        print(f"Please add .docx, .pdf, or .txt files to check.")
        sys.exit(0)

    print(f"\nFiles available in {FILES_DIR.name}/:\n")
    for i, f in enumerate(files, 1):
        print(f"  {i}. {f.name}")

    # Get user selection
    print()
    while True:
        try:
            choice = input("Select file number (or 'q' to quit): ").strip()
            if choice.lower() == 'q':
                print("Goodbye!")
                sys.exit(0)

            choice_idx = int(choice) - 1
            if 0 <= choice_idx < len(files):
                selected_file = files[choice_idx]
                break
            else:
                print("Invalid selection. Try again.")
        except ValueError:
            print("Please enter a number.")

    # Get target percentage
    while True:
        try:
            target = input("Enter target AI percentage (e.g., 20): ").strip()
            target_percentage = float(target)
            if 0 <= target_percentage <= 100:
                break
            else:
                print("Please enter a value between 0 and 100.")
        except ValueError:
            print("Please enter a valid number.")

    print(f"\n{'=' * 50}")
    print(f"Processing: {selected_file.name}")
    print(f"Target: Below {target_percentage}% AI content")
    print(f"{'=' * 50}")

    # Initial scan
    print("\n[1] Scanning original file...")
    result = detect_ai_in_file(selected_file)
    display_results(result, selected_file.name)

    current_percentage = parse_ai_percentage(result)

    if current_percentage < 0:
        print("Error: Could not get AI percentage from scan.")
        sys.exit(1)

    if current_percentage <= target_percentage:
        print(f"\nFile is already below target ({current_percentage}% <= {target_percentage}%)")
        print("No humanisation needed!")
        sys.exit(0)

    # Extract text for humanisation
    print("\n[2] Extracting text from file...")
    current_text = extract_text_from_file(selected_file)

    if not current_text.strip():
        print("Error: Could not extract text from file.")
        sys.exit(1)

    print(f"  Extracted {len(current_text.split())} words")

    # Humanisation loop
    version = 1
    max_iterations = 10  # Safety limit
    current_file = selected_file

    while current_percentage > target_percentage and version <= max_iterations:
        print(f"\n[{version + 2}] Humanising text (attempt {version})...")

        # Humanise the text
        humanised_text = humanise_text(current_text)

        # Save to new file
        output_file = save_text_to_file(humanised_text, selected_file, version)
        print(f"  Saved to: {output_file.name}")

        # Re-scan
        print(f"\n[{version + 3}] Re-scanning humanised file...")

        # Use text detection for re-scan (faster and works with our output)
        result = detect_ai_in_text(humanised_text)
        display_results(result, output_file.name)

        current_percentage = parse_ai_percentage(result)

        if current_percentage < 0:
            print("Error: Could not get AI percentage from re-scan.")
            break

        if current_percentage <= target_percentage:
            print(f"\n{'=' * 50}")
            print(f"  SUCCESS!")
            print(f"  Final AI percentage: {current_percentage}%")
            print(f"  Target was: {target_percentage}%")
            print(f"  Output file: {output_file}")
            print(f"{'=' * 50}")
            break

        # Prepare for next iteration
        current_text = humanised_text
        version += 1

        if version > max_iterations:
            print(f"\nReached maximum iterations ({max_iterations}).")
            print(f"Current AI percentage: {current_percentage}%")
            print(f"Latest file: {output_file}")

    print("\nDone!")


if __name__ == "__main__":
    main()
